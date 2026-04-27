"""
Letter Generator Backend — FastAPI
POST /api/letter/generate-zip

Reads SUMMARY sheet from each MERGED Excel file,
injects data into MASTER LETTER .docx template,
returns a ZIP of all generated Word files.
"""

from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

import io
import copy
import zipfile
from pathlib import Path
from typing import List

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Asset paths (signature + company stamp) ──────────────────────────────────
ASSETS_DIR = Path(__file__).parent / "assets"
SIGNATURE_PNG = ASSETS_DIR / "signature.png"
STAMP_PNG     = ASSETS_DIR / "stamp.png"

# ── Country code → full name map ──────────────────────────────────────────────
COUNTRY_MAP = {
    "CN": "CHINA", "DE": "GERMANY", "KR": "KOREA", "CZ": "CZECH",
    "ES": "SPAIN", "FR": "FRANCE", "PL": "POLAND", "RO": "ROMANIA",
    "SK": "SLOVAKIA", "TR": "TURKEY", "VN": "VIETNAM", "TH": "THAILAND",
    "US": "USA", "JP": "JAPAN", "IT": "ITALY", "GB": "UK",
    "NL": "NETHERLANDS", "SE": "SWEDEN", "AT": "AUSTRIA", "BE": "BELGIUM",
    "HU": "HUNGARY", "PT": "PORTUGAL", "DK": "DENMARK", "FI": "FINLAND",
    "NO": "NORWAY", "CH": "SWITZERLAND", "IN": "INDIA", "BR": "BRAZIL",
    "MX": "MEXICO", "AU": "AUSTRALIA", "MY": "MALAYSIA", "TW": "TAIWAN",
    "SG": "SINGAPORE", "ID": "INDONESIA", "PH": "PHILIPPINES",
}

REQUIRED_COLS = [
    "Country",
    "Total Invoiced Qty",
    "Total Net Weight",
    "Total Net Value (USD)",
]

# Optional columns — if missing in the file, treated as empty (does not fail validation).
# Older MERGED files don't have this column; we still want letters to generate.
OPTIONAL_COLS = [
    "Total GW.",
]

app = FastAPI(title="Letter Generator API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Helpers ───────────────────────────────────────────────────────────────────

def resolve_country(code: str) -> str:
    """Return 'MADE IN {COUNTRY NAME}' from a code or full name."""
    c = str(code).strip().upper()
    name = COUNTRY_MAP.get(c, c)          # fallback: use value as-is if not in map
    return f"MADE IN {name}"


def fmt_qty(v) -> str:
    try:
        return str(int(float(v)))
    except Exception:
        return str(v)


def fmt_weight(v) -> str:
    try:
        return f"{float(v):.3f}"
    except Exception:
        return str(v)


def fmt_usd(v) -> str:
    try:
        return f"{float(v):,.2f}"
    except Exception:
        return str(v)


def parse_summary(excel_bytes: bytes, filename: str) -> dict:
    """
    Read SUMMARY sheet from excel_bytes.
    Returns:
      {
        "country_rows": [ {origin, qty, nw, gw, fob} ... ],
        "total_row":    {origin, qty, nw, gw, fob} | None,
        "country_list": "CZECH, GERMANY, ..."
      }
    Raises ValueError with a clear message on validation failure.
    """
    try:
        xl = pd.ExcelFile(io.BytesIO(excel_bytes))
    except Exception as e:
        raise ValueError(f"อ่าน Excel ไม่ได้: {e}")

    if "SUMMARY" not in xl.sheet_names:
        # Case-insensitive fallback
        found = next((s for s in xl.sheet_names if s.strip().upper() == "SUMMARY"), None)
        if not found:
            raise ValueError(
                f"ไม่พบ Sheet 'SUMMARY' ใน {filename} "
                f"(sheets ที่พบ: {', '.join(xl.sheet_names)})"
            )
        sheet_name = found
    else:
        sheet_name = "SUMMARY"

    # Read with no header first so we can auto-find the real header row.
    # Real files often have a title row at the top (e.g. "SUMMARY — Country of Origin Report"),
    # which makes pandas's default header=0 grab the title instead of real column names.
    raw = pd.read_excel(io.BytesIO(excel_bytes), sheet_name=sheet_name, header=None, dtype=object)

    # Normalize helper — tolerant to spacing, casing, and trailing punctuation.
    # Real files in the wild use both "Total GW" (no dot) and "Total GW." — same column.
    # Same for "Total Net Weight" vs "Total Net Weight." etc.
    def _norm(s):
        cleaned = " ".join(str(s).strip().split()).upper()
        # Drop trailing punctuation (dots, commas) so "GW" == "GW." etc.
        return cleaned.rstrip(" .,:;")

    header_row_idx = None
    scan_limit = min(len(raw), 20)  # scan up to 20 rows (safer than 10)
    scanned_rows = []  # for error reporting

    for i in range(scan_limit):
        row_vals_raw = [str(v).strip() for v in raw.iloc[i].values if v is not None and str(v).strip()]
        scanned_rows.append((i, row_vals_raw))
        if not row_vals_raw:
            continue

        # Normalize both sides for comparison (handles extra spaces, case differences)
        row_norm = {_norm(s) for s in row_vals_raw}
        required_norm = {_norm(r) for r in REQUIRED_COLS}
        hits = len(row_norm & required_norm)

        if hits >= 3:
            header_row_idx = i
            break
        # Alternative: "Country" + any one "Total ..." column
        if _norm("Country") in row_norm and any("TOTAL" in s for s in row_norm):
            header_row_idx = i
            break

    if header_row_idx is None:
        # Build debug message showing what we actually found
        debug_lines = []
        for idx, vals in scanned_rows[:10]:  # show first 10 rows only
            preview = " | ".join(vals[:8])  # first 8 cells max
            debug_lines.append(f"  Row {idx}: {preview}")
        debug = "\n".join(debug_lines) if debug_lines else "(ไฟล์ว่าง)"
        raise ValueError(
            f"{filename} — หา header row ไม่เจอใน SUMMARY sheet\n"
            f"ต้องมีคอลัมน์เหล่านี้: {', '.join(REQUIRED_COLS)}\n"
            f"แต่ใน 10 แถวแรกของ sheet มีข้อมูลดังนี้:\n{debug}"
        )

    df = pd.read_excel(
        io.BytesIO(excel_bytes), sheet_name=sheet_name,
        header=header_row_idx, dtype=object
    )
    df.columns = [str(c).strip() for c in df.columns]

    # Validate required columns using normalized matching, and rename
    # user-file columns to canonical names so the rest of the code can
    # access them by the spec-defined names.
    # Both REQUIRED_COLS and OPTIONAL_COLS get renamed (so "Total GW" → "Total GW.")
    canonical_by_norm = {_norm(c): c for c in (REQUIRED_COLS + OPTIONAL_COLS)}
    rename_map = {}
    for col in df.columns:
        n = _norm(col)
        if n in canonical_by_norm and canonical_by_norm[n] != col:
            rename_map[col] = canonical_by_norm[n]
    if rename_map:
        df = df.rename(columns=rename_map)

    # Add empty optional columns if not present (older files don't have "Total GW.")
    for opt in OPTIONAL_COLS:
        if opt not in df.columns:
            df[opt] = ""

    # Now check for missing REQUIRED columns (normalized)
    present_norm = {_norm(c) for c in df.columns}
    required_norm = {_norm(c) for c in REQUIRED_COLS}
    missing_norm = required_norm - present_norm
    if missing_norm:
        # Map back to canonical names for the error message
        missing = [c for c in REQUIRED_COLS if _norm(c) in missing_norm]
        raise ValueError(
            f"{filename} — คอลัมน์ที่ขาด: {', '.join(missing)} "
            f"(พบคอลัมน์: {', '.join(df.columns)})"
        )

    # Drop fully empty rows
    df = df.dropna(how="all")

    # Normalize Country column to string for comparison
    country_series = df["Country"].astype(str).str.strip()

    # Separate GRAND TOTAL row from data rows
    is_total = country_series.str.upper() == "GRAND TOTAL"
    total_df = df[is_total]
    data_df  = df[~is_total & country_series.ne("") & country_series.str.upper().ne("NAN")]

    if data_df.empty:
        raise ValueError(f"{filename} — ไม่พบแถวข้อมูลใน SUMMARY")
    if total_df.empty:
        raise ValueError(f"{filename} — ไม่พบแถว 'GRAND TOTAL' ใน SUMMARY")

    country_rows = []
    seen = set()
    country_list_parts = []
    for _, row in data_df.iterrows():
        country_code = str(row["Country"]).strip()
        if not country_code:
            continue
        origin = resolve_country(country_code)
        country_rows.append({
            "origin": origin,
            "qty":    fmt_qty(row["Total Invoiced Qty"]),
            "nw":     fmt_weight(row["Total Net Weight"]),
            "gw":     fmt_weight(row["Total GW."]),
            "fob":    fmt_usd(row["Total Net Value (USD)"]),
        })
        # Build country list (unique, preserving order, using full name)
        full = COUNTRY_MAP.get(country_code.upper(), country_code.upper())
        if full not in seen:
            seen.add(full)
            country_list_parts.append(full)

    total_row_raw = total_df.iloc[0]
    total_row = {
        "origin": "TOTAL",
        "qty":    fmt_qty(total_row_raw["Total Invoiced Qty"]),
        "nw":     fmt_weight(total_row_raw["Total Net Weight"]),
        "gw":     fmt_weight(total_row_raw["Total GW."]),
        "fob":    fmt_usd(total_row_raw["Total Net Value (USD)"]),
    }

    return {
        "country_rows": country_rows,
        "total_row":    total_row,
        "country_list": ", ".join(country_list_parts),
    }


# ── Word manipulation ─────────────────────────────────────────────────────────

def _replace_text_in_run(run, old: str, new: str):
    if old in run.text:
        run.text = run.text.replace(old, new)


def replace_placeholders(doc: Document, mapping: dict):
    """
    Replace all {{KEY}} placeholders in the document.
    Handles both single-run and split-run cases.
    """
    for para in doc.paragraphs:
        _replace_para_placeholders(para, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_para_placeholders(para, mapping)

    # Headers and footers
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer]:
            if hf is not None:
                for para in hf.paragraphs:
                    _replace_para_placeholders(para, mapping)
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _replace_para_placeholders(para, mapping)


def _replace_para_placeholders(para, mapping: dict):
    """
    Robust placeholder replacement that handles runs split by Word's XML engine.
    Merges all run text, replaces, then rewrites back into first run.
    """
    full = "".join(r.text for r in para.runs)
    if "{{" not in full:
        return
    for key, val in mapping.items():
        full = full.replace(f"{{{{{key}}}}}", str(val))
    # Write merged text back — put all into first run, clear the rest
    if para.runs:
        para.runs[0].text = full
        for r in para.runs[1:]:
            r.text = ""


def find_origin_table(doc: Document):
    """
    Find the table whose header row contains the columns:
    ORIGIN, QTY, N.W., G.W., FOB USD (in any cell, case-insensitive).
    Returns the table or None.
    """
    target_headers = {"ORIGIN", "QTY", "N.W.", "G.W.", "FOB USD"}
    for table in doc.tables:
        if not table.rows:
            continue
        header_row = table.rows[0]
        cell_texts = {c.text.strip().upper() for c in header_row.cells}
        if target_headers.issubset(cell_texts):
            return table
    return None


def _copy_row_format(template_row, new_row):
    """Copy XML formatting (trPr) from template_row to new_row."""
    try:
        src_trpr = template_row._tr.find(qn("w:trPr"))
        dst_trpr = new_row._tr.find(qn("w:trPr"))
        if src_trpr is not None:
            if dst_trpr is not None:
                new_row._tr.remove(dst_trpr)
            new_row._tr.insert(0, copy.deepcopy(src_trpr))
    except Exception:
        pass


def _set_cell_text(cell, text: str, template_cell=None):
    """Set cell text while preserving formatting from template_cell."""
    # Clear existing paragraphs content
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""

    if template_cell is not None:
        # Copy cell properties (width, shading, borders)
        try:
            src_tcpr = template_cell._tc.find(qn("w:tcPr"))
            dst_tcpr = cell._tc.find(qn("w:tcPr"))
            if src_tcpr is not None:
                if dst_tcpr is not None:
                    cell._tc.remove(dst_tcpr)
                cell._tc.insert(0, copy.deepcopy(src_tcpr))
        except Exception:
            pass

        # Copy paragraph formatting and set text
        src_paras = template_cell.paragraphs
        dst_paras = cell.paragraphs
        for i, src_para in enumerate(src_paras):
            if i < len(dst_paras):
                dst_para = dst_paras[i]
            else:
                dst_para = cell.add_paragraph()
            # Copy paragraph properties
            try:
                src_ppr = src_para._p.find(qn("w:pPr"))
                dst_ppr = dst_para._p.find(qn("w:pPr"))
                if src_ppr is not None:
                    if dst_ppr is not None:
                        dst_para._p.remove(dst_ppr)
                    dst_para._p.insert(0, copy.deepcopy(src_ppr))
            except Exception:
                pass
            # Set text via run (copy run format from first run of src_para)
            if src_para.runs:
                src_run = src_para.runs[0]
                if dst_para.runs:
                    run = dst_para.runs[0]
                else:
                    run = dst_para.add_run()
                # Copy run properties
                try:
                    src_rpr = src_run._r.find(qn("w:rPr"))
                    dst_rpr = run._r.find(qn("w:rPr"))
                    if src_rpr is not None:
                        if dst_rpr is not None:
                            run._r.remove(dst_rpr)
                        run._r.insert(0, copy.deepcopy(src_rpr))
                except Exception:
                    pass
                run.text = str(i == 0 and text or "")  # only first para gets text
            else:
                if dst_para.runs:
                    dst_para.runs[0].text = text if i == 0 else ""
                else:
                    dst_para.add_run(text if i == 0 else "")
    else:
        # Fallback: just write text
        if cell.paragraphs:
            para = cell.paragraphs[0]
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)
        else:
            cell.add_paragraph(text)


def _get_col_index(header_row, col_name: str) -> int:
    """Return index of cell matching col_name (case-insensitive)."""
    for i, cell in enumerate(header_row.cells):
        if cell.text.strip().upper() == col_name.upper():
            return i
    return -1


def _replace_in_paragraph_text(para, old: str, new: str):
    """
    Replace `old` with `new` inside a paragraph's text, preserving the first
    run's formatting. Works when the text is split across multiple runs by Word.
    """
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    updated = full.replace(old, new)
    if para.runs:
        para.runs[0].text = updated
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(updated)
    return True


def _iter_all_paragraphs(doc: Document):
    """Yield every paragraph in body, tables, headers, footers."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer]:
            if hf is None:
                continue
            for para in hf.paragraphs:
                yield para
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para


def _apply_replacements_safely(doc: Document, replacements: list):
    """
    Apply replacements while ensuring no replacement's output is re-matched
    by a later replacement. Each paragraph is processed once, applying all
    rules sequentially using unique placeholder tokens to mark already-replaced
    regions.
    """
    import re as _re
    for para in _iter_all_paragraphs(doc):
        full = "".join(r.text for r in para.runs)
        if not full.strip():
            continue
        original = full
        # Use sentinel tokens to mark replacements; restore at end.
        sentinels = []
        def stash(val):
            idx = len(sentinels)
            sentinels.append(val)
            return f"\x00SENT{idx}\x00"

        new_text = full
        changed = False
        for pat, repl in replacements:
            if hasattr(pat, "sub"):
                def _wrap(m, _pat=pat, _repl=repl):
                    if callable(_repl):
                        out = _repl(m)
                    else:
                        out = m.expand(_repl)
                    return stash(out)
                nt = pat.sub(_wrap, new_text)
            else:
                if pat in new_text:
                    nt = new_text.replace(pat, stash(repl))
                else:
                    nt = new_text
            if nt != new_text:
                new_text = nt
                changed = True
        # Restore sentinels
        def _restore(m):
            return sentinels[int(m.group(1))]
        new_text = _re.sub(r"\x00SENT(\d+)\x00", _restore, new_text)
        if changed and para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""


def apply_text_replacements(doc: Document, replacements: list):
    """
    replacements: list of (pattern, replacement) tuples.
    pattern can be a plain string (exact substring) or a compiled regex.
    Applied to every paragraph's concatenated text. First replacement that
    matches wins for that paragraph, so put more-specific patterns first.
    """
    for para in _iter_all_paragraphs(doc):
        full = "".join(r.text for r in para.runs)
        if not full.strip():
            continue
        changed = False
        new_text = full
        for pat, repl in replacements:
            if hasattr(pat, "sub"):  # compiled regex
                nt = pat.sub(repl, new_text)
                if nt != new_text:
                    new_text = nt
                    changed = True
            else:
                if pat in new_text:
                    new_text = new_text.replace(pat, repl)
                    changed = True
        if changed and para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""


def inject_table_rows(doc: Document, country_rows: list, total_row: dict):
    """
    Find the ORIGIN table, keep header row,
    remove all existing data rows,
    inject new country rows + TOTAL row,
    copying cell/row formatting from the original first data row.
    """
    table = find_origin_table(doc)
    if table is None:
        # Provide helpful diagnostic about what was actually found
        num_tables = len(doc.tables)
        if num_tables == 0:
            hint = (
                "ไฟล์ MASTER นี้ไม่มีตารางเลย — "
                "น่าจะเป็นไฟล์หัวจดหมาย (letterhead) ไม่ใช่ template เต็ม "
                "กรุณาใช้ไฟล์ template ที่มีตาราง ORIGIN เช่น 'จดหมาย_SCHAEFFLER.docx'"
            )
        else:
            # List what headers we DID find
            found_headers = []
            for i, t in enumerate(doc.tables):
                if t.rows:
                    cells = [c.text.strip() for c in t.rows[0].cells]
                    found_headers.append(f"Table {i}: {cells}")
            hint = (
                f"ไฟล์ MASTER มี {num_tables} ตาราง แต่ไม่มีตารางไหนที่มี header "
                f"ORIGIN | QTY | N.W. | G.W. | FOB USD\n"
                f"ตารางที่พบ:\n  " + "\n  ".join(found_headers)
            )
        raise ValueError(f"ไม่พบตาราง ORIGIN ใน MASTER — {hint}")

    header_row  = table.rows[0]
    # Identify column indices from header
    col_origin  = _get_col_index(header_row, "ORIGIN")
    col_qty     = _get_col_index(header_row, "QTY")
    col_nw      = _get_col_index(header_row, "N.W.")
    col_gw      = _get_col_index(header_row, "G.W.")
    col_fob     = _get_col_index(header_row, "FOB USD")

    # Keep template (first data row) for formatting reference, then delete all data rows
    template_row = table.rows[1] if len(table.rows) > 1 else None
    template_cells = [c for c in template_row.cells] if template_row else None

    # Remove all rows except header
    tbl_elem = table._tbl
    rows_to_remove = list(table.rows)[1:]
    for row in rows_to_remove:
        tbl_elem.remove(row._tr)

    def _force_cell_style(cell, bold_all: bool):
        """
        Ensure every run in the cell has: color=black, and optionally bold=True.
        Keeps font name/size from template — only adjusts color and bold weight.
        """
        from docx.shared import RGBColor
        BLACK = RGBColor(0x00, 0x00, 0x00)
        for para in cell.paragraphs:
            for run in para.runs:
                # Always force black text (user request)
                run.font.color.rgb = BLACK
                if bold_all:
                    run.bold = True

    def add_data_row(data: dict, is_total: bool = False):
        """Add a new row to the table using the template row's format."""
        new_tr = copy.deepcopy(template_row._tr) if template_row else OxmlElement("w:tr")
        tbl_elem.append(new_tr)
        # Find the new row object
        new_row = table.rows[-1]

        # Map column index → value
        col_map = {
            col_origin: data.get("origin", ""),
            col_qty:    data.get("qty",    ""),
            col_nw:     data.get("nw",     ""),
            col_gw:     data.get("gw",     ""),
            col_fob:    data.get("fob",    ""),
        }

        for ci, cell in enumerate(new_row.cells):
            val = col_map.get(ci, "")
            tmpl_cell = template_cells[ci] if template_cells and ci < len(template_cells) else None
            _set_cell_text(cell, val, tmpl_cell)
            # Force black color; for TOTAL row, make every cell bold
            _force_cell_style(cell, bold_all=is_total)

    # Inject country rows
    for row_data in country_rows:
        add_data_row(row_data, is_total=False)

    # Inject TOTAL row last — all cells bold
    add_data_row(total_row, is_total=True)


def inject_signature_and_stamp(doc: Document) -> bool:
    """
    Insert signature image + company stamp into the empty paragraph(s) that sit
    between "ขอแสดงความนับถือ" and the signer name line.

    Layout produced (single line, side by side):
        [signature]   [stamp]

    Returns True if successfully injected, False if anchors weren't found
    or assets are missing (silent — letter still generates without images).
    """
    if not SIGNATURE_PNG.exists() and not STAMP_PNG.exists():
        return False

    # 1) Find the "ขอแสดงความนับถือ" paragraph (anchor)
    paragraphs = doc.paragraphs
    anchor_idx = None
    for i, p in enumerate(paragraphs):
        if "ขอแสดงความนับถือ" in p.text:
            anchor_idx = i
            break
    if anchor_idx is None:
        return False

    # 2) Find the signer name line below (starts with นาย/นาง/นางสาว) — to know how many
    #    blank lines we have between them. We'll target the blank line right after the
    #    anchor (typically anchor+1 or anchor+2).
    target_para = None
    # Prefer first empty/whitespace-only paragraph after the anchor
    for j in range(anchor_idx + 1, min(anchor_idx + 6, len(paragraphs))):
        if not paragraphs[j].text.strip():
            target_para = paragraphs[j]
            break

    # Fallback: just use the next paragraph
    if target_para is None and anchor_idx + 1 < len(paragraphs):
        target_para = paragraphs[anchor_idx + 1]

    if target_para is None:
        return False

    # 3) Clear any existing runs in this paragraph and insert our images
    for run in list(target_para.runs):
        run._element.getparent().remove(run._element)

    target_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Indent so the images appear under the signer area (roughly center-left)
    try:
        target_para.paragraph_format.left_indent = Cm(7.0)
    except Exception:
        pass

    # Insert signature first, then a small spacer, then stamp
    if SIGNATURE_PNG.exists():
        run_sig = target_para.add_run()
        try:
            run_sig.add_picture(str(SIGNATURE_PNG), width=Cm(3.0))
        except Exception:
            pass

    if STAMP_PNG.exists():
        run_stamp = target_para.add_run("  ")  # small gap
        try:
            run_stamp.add_picture(str(STAMP_PNG), width=Cm(2.8))
        except Exception:
            pass

    return True


def generate_letter(
    excel_bytes: bytes,
    master_bytes: bytes,
    filename: str,
    letter_date: str,
    invoice_no: str,
    invoice_date: str,
    destination_country: str,
    signer_name: str,
    signer_position: str,
) -> bytes:
    """
    Full pipeline for one MERGED file:
    1. Parse SUMMARY
    2. Open master template
    3. Replace {{placeholders}} (if any)
    4. Apply pattern-based replacements for templates without placeholders
    5. Inject table
    6. Return .docx bytes
    """
    import re

    # Step 1 — parse Excel
    summary = parse_summary(excel_bytes, filename)

    # Step 2 — open fresh copy of master template
    doc = Document(io.BytesIO(master_bytes))

    # Step 3 — replace {{...}} placeholders (no-op if template has none)
    placeholders = {
        "DATE":                letter_date,
        "INVOICE_NO":          invoice_no,
        "INVOICE_DATE":        invoice_date,
        "GW_TOTAL":            summary["total_row"]["gw"],
        "DESTINATION_COUNTRY": destination_country,
        "COUNTRY_LIST":        summary["country_list"],
        "SIGNER_NAME":         signer_name,
        "SIGNER_POSITION":     signer_position,
    }
    replace_placeholders(doc, placeholders)

    # Step 4 — Pattern-based replacements for templates without placeholders.
    # Uses Thai patterns that match the actual "จดหมาย_SCHAEFFLER.docx" template.
    gw_total  = summary["total_row"]["gw"]
    country_list = summary["country_list"]

    # Build replacements list (order matters — MOST-SPECIFIC FIRST).
    # The invoice block MUST run before letter_date, because P5 contains
    # "วันที่" which a generic date rule would otherwise catch.
    replacements = []

    # 1. เลขที่อินวอยซ์ <INV_NO> วันที่ <INV_DATE> — MUST run first.
    # Collapse the long padding whitespace to single spaces for a clean look:
    #   "เลขที่อินวอยซ์                 113979693  วันที่              21/04/2569"
    # becomes:
    #   "เลขที่อินวอยซ์ 113979693 วันที่ 21/04/2569"
    #
    # Problem to avoid: the raw template has EMPTY values between labels, i.e.
    #   "เลขที่อินวอยซ์<spaces>วันที่<spaces>น้ำหนัก..."
    # A naive \S+ would match "น้ำหนัก" itself as the date value and eat it.
    # Solution: the "value" chars are restricted to [\w/.,\-—:] — no Thai letters,
    # no spaces — so the regex stops at the first "real" word like น้ำหนัก.
    if invoice_no or invoice_date:
        VALCHR = r"[A-Za-z0-9/.,\-_]"
        def _inv_repl(m):
            inv_part  = f"เลขที่อินวอยซ์ {invoice_no}" if invoice_no else "เลขที่อินวอยซ์"
            date_part = f"วันที่ {invoice_date}" if invoice_date else "วันที่"
            return f"{inv_part} {date_part} "
        replacements.append((
            re.compile(
                rf"เลขที่อินวอยซ์[ \t\u00A0]+(?:{VALCHR}+[ \t\u00A0]+)?วันที่[ \t\u00A0]+(?:{VALCHR}+[ \t\u00A0]+)?",
                re.UNICODE
            ),
            _inv_repl
        ))

    # 2. Letter Date — matches ONLY "วันที่ dd/mm/yyyy" standalone date line.
    # The explicit dd/mm/yyyy anchor prevents matching P5's empty placeholder.
    if letter_date:
        replacements.append((
            re.compile(r"(วันที่\s+)\d{1,2}/\d{1,2}/\d{2,4}", re.UNICODE),
            rf"\g<1>{letter_date}"
        ))

    # 3. Destination country ("ประเทศ VIETNAM")
    if destination_country:
        replacements.append((
            re.compile(
                r"(ได้ส่งสินค้าไปยังประเทศ\s+)([A-Z][A-Z\s]+?)(\s+เลขที่อินวอยซ์)",
                re.UNICODE
            ),
            rf"\g<1>{destination_country}\g<3>"
        ))

    # 4. น้ำหนัก G.W <value> KGM — ensure single space before KGM
    if gw_total:
        replacements.append((
            re.compile(r"น้ำหนัก\s+G\.?\s*W\s*[\d,.\s]*\s*KGM", re.UNICODE | re.IGNORECASE),
            f"น้ำหนัก G.W {gw_total} KGM"
        ))

    # 5. Country list — "CZECH, GERMANY, ... และ CHINA"
    if country_list:
        parts = [p.strip() for p in country_list.split(",") if p.strip()]
        if len(parts) >= 2:
            thai_list = ", ".join(parts[:-1]) + " และ " + parts[-1]
        else:
            thai_list = country_list
        replacements.append((
            re.compile(
                r"(ทีมีแหล่งผลิตจากประเทศ\s+).+?(\s+หากเก)",
                re.UNICODE
            ),
            rf"\g<1>{thai_list}\g<2>"
        ))

    # 6. Signer name — ONLY whole-line Thai name paragraph (^...$)
    if signer_name:
        replacements.append((
            re.compile(r"^(นาย|นาง|นางสาว)\s+[\u0E00-\u0E7F\s]+$", re.UNICODE),
            signer_name
        ))

    # 7. Signer position
    if signer_position:
        replacements.append(("Schaeffler Manufacturing (Thailand) Co., Ltd.", signer_position))

    # Apply using a two-phase token system to avoid re-matching values
    # already substituted (e.g. letter_date regex accidentally matching
    # the invoice_date we just inserted).
    if replacements:
        _apply_replacements_safely(doc, replacements)

    # Step 5 — inject table
    inject_table_rows(doc, summary["country_rows"], summary["total_row"])

    # Step 6 — inject signature + company stamp into the empty area below "ขอแสดงความนับถือ"
    inject_signature_and_stamp(doc)

    # Step 7 — save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── API endpoint ──────────────────────────────────────────────────────────────

@app.post("/api/letter/generate-zip")
async def generate_zip(
    merged_files:        List[UploadFile] = File(...),
    master_file:         UploadFile       = File(...),
    letter_date:         str = Form(""),
    invoice_no:          str = Form(""),
    invoice_date:        str = Form(""),
    destination_country: str = Form(""),
    signer_name:         str = Form(""),
    signer_position:     str = Form(""),
):
    # Read master template bytes once
    master_bytes = await master_file.read()
    if not master_file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Master file ต้องเป็น .docx")

    zip_buf = io.BytesIO()
    errors  = []
    success = 0

    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, mf in enumerate(merged_files, start=1):
            excel_bytes = await mf.read()
            stem = mf.filename.rsplit(".", 1)[0]
            out_name = f"CO_REQUEST_{i:03d}_{stem}.docx"

            try:
                docx_bytes = generate_letter(
                    excel_bytes=excel_bytes,
                    master_bytes=master_bytes,
                    filename=mf.filename,
                    letter_date=letter_date,
                    invoice_no=invoice_no,
                    invoice_date=invoice_date,
                    destination_country=destination_country,
                    signer_name=signer_name,
                    signer_position=signer_position,
                )
                zf.writestr(out_name, docx_bytes)
                success += 1
            except ValueError as e:
                errors.append(f"{mf.filename}: {e}")
            except Exception as e:
                errors.append(f"{mf.filename}: unexpected error — {e}")

    if success == 0:
        detail = "ไม่มีไฟล์ที่สำเร็จ\n" + "\n".join(errors)
        raise HTTPException(422, detail=detail)

    # If some failed, add an error log inside the ZIP
    if errors:
        err_text = "ERRORS:\n" + "\n".join(errors)
        with zipfile.ZipFile(zip_buf, "a") as zf:
            zf.writestr("ERRORS.txt", err_text)

    zip_buf.seek(0)
    return StreamingResponse(
        zip_buf,
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=CO_REQUEST_BATCH.zip"},
    )
